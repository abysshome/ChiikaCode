import uvicorn
from fastapi import FastAPI, HTTPException
from generate_code_api import (RequestData, generate_code_structure,
                               supported_langs)

app = FastAPI()

import csv
import os

import openpyxl
import uvicorn
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document
from langchain.schema.output_parser import StrOutputParser
from langchain.schema.runnable import RunnablePassthrough
from langchain.text_splitter import (CharacterTextSplitter,
                                     RecursiveCharacterTextSplitter)
from langchain_community.chat_models import ChatOllama
from langchain_community.document_loaders import (CSVLoader, PyPDFLoader,
                                                  PythonLoader, TextLoader)
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import Chroma
from openpyxl import load_workbook
from pdfminer.high_level import extract_text
from pydantic import BaseModel


class BaseLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        raise NotImplementedError("This method should be overridden by subclasses")

class PythonLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class PyPDFLoader(BaseLoader):
    def load(self):
        return extract_text(self.file_path)

class TextLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class CSVLoader(BaseLoader):
    def load(self):
        with open(self.file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            return list(reader)

class DocxLoader(BaseLoader):
    def load(self):
        doc = DocxDocument(self.file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

class XlsxLoader(BaseLoader):
    def load(self):
        workbook = load_workbook(filename=self.file_path)
        sheet = workbook.active
        data = []
        for row in sheet.iter_rows(values_only=True):
            data.append(row)
        return data

class CLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class JavaLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class CPPLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

# 定义一个请求模型，用于接收文件路径
class FilePathRequest(BaseModel):
    file_path: str

# 数据加载
def load_file(file_path):
    if not os.path.exists(file_path):
        raise HTTPException(status_code=400, detail="文件路径不存在")

    if file_path.endswith('.py'):
        loader = PythonLoader(file_path)
    elif file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path)
    elif file_path.endswith('.csv'):
        loader = CSVLoader(file_path)
    elif file_path.endswith('.docx'):
        loader = DocxLoader(file_path)
    elif file_path.endswith('.xlsx'):
        loader = XlsxLoader(file_path)
    elif file_path.endswith('.c'):
        loader = CLoader(file_path)
    elif file_path.endswith('.java'):
        loader = JavaLoader(file_path)
    elif file_path.endswith('.cpp'):
        loader = CPPLoader(file_path)
    else:
        print("不支持的文件类型")
        return None
    
    documents = loader.load()
    return documents

# 文档分割
def split_documents(documents):
    # 将字符串包装成 Document 对象
    if isinstance(documents, str):
        documents = [Document(page_content=documents)]
    elif isinstance(documents, list) and all(isinstance(item, str) for item in documents):
        documents = [Document(page_content=item) for item in documents]
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=65536, chunk_overlap=10)
    print("文档分割中...")
    documents = text_splitter.split_documents(documents)
    print("文档分割完成")
    return documents

# 嵌入向量获取
def get_embedding():
    model_name = 'moka-ai/m3e-base'
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': True}
    embedding = HuggingFaceBgeEmbeddings(
        model_name=model_name, 
        model_kwargs=model_kwargs, 
        encode_kwargs=encode_kwargs
    )
    return embedding

# 构建向量数据库
def build_vector_db(documents, embedding):
    persist_directory = 'db'
    db = Chroma.from_documents(
        documents, 
        embedding, 
        persist_directory=persist_directory
    )
    return db

# 获取向量数据库
def get_vector_db(embedding):
    persist_directory = 'db'
    db = Chroma(persist_directory=persist_directory
                , embedding_function=embedding)
    return db

# 构建检索器
def build_retriever(db):
    retriever = db.as_retriever()
    return retriever

# 构建RAG链
def build_rag_chain(retriever):
    template = """
        根据context详细解释有关question的内容，并给出答案。
        Question: {question} Context: {context} Answer: """
    prompt = ChatPromptTemplate.from_template(template)
    llm = ChatOllama(model='llama3.2:latest')
    rag_chain = (
        {"context": retriever, "question": RunnablePassthrough()} |
        prompt |
        llm |
        StrOutputParser()
    )
    return rag_chain

# 处理文件夹
def handle_folder(folder_path):
    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)
        if os.path.isdir(file_path):
            handle_folder(file_path)
        else:
            print(file_path)
            if file_path.endswith('.py') or file_path.endswith('.pdf') or file_path.endswith('.txt') or file_path.endswith('.csv'):
                documents = load_file(file_path)
                documents = split_documents(documents)
                build_vector_db(documents, embedding)
                
# 全局变量，用于保存 RAG 链
rag_chain = None

# 上传文件路径并初始化 RAG 链
@app.post("/upload_file_path")
async def upload_file_path(request: FilePathRequest):
    print(request)
    global rag_chain
    global embedding
    embedding = get_embedding()
    try:
        # 如果是文件夹
        if os.path.isdir(request.file_path):
            handle_folder(request.file_path)
        # 如果是文件
        else:
            documents = load_file(request.file_path)# 加载文件
            documents = split_documents(documents)
            build_vector_db(documents, embedding)
        # 获取向量数据库
        db=get_vector_db(embedding)
        retriever = build_retriever(db)
        rag_chain = build_rag_chain(retriever)
        return {"message": "文件路径上传成功，RAG链已初始化"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"内部服务器错误: {str(e)}")

# 定义提问模型
class QuestionRequest(BaseModel):
    question: str

# 提问并获取答案
@app.post("/rag_ask")
async def ask_question(request: QuestionRequest):
    print(request)
    global rag_chain
    if not rag_chain:
        raise HTTPException(status_code=400, detail="RAG链尚未初始化，请先上传文件路径")

    question = request.question
    answer = ""
    async for chunk in rag_chain.astream(question):
        answer += chunk
    print(answer)
    return {"answer": answer}

# 项目级代码生成
@app.post("/generate_project")
async def generate_project(data: RequestData):
    print(data)
    # 检查语言是否支持
    if data.language not in supported_langs:
        raise HTTPException(status_code=400, detail=f"语言 '{data.language}' 不受支持")
    print("语言支持")
    # 调用生成代码结构的函数
    try:
        result_json = generate_code_structure(data.question, data.language)
        print("generate_code_structure 成功")
        return {"status": "success", "message": "项目代码生成完成", "result": result_json}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成代码时出错: {str(e)}")
    
from executable_code_api import CodeRequest, CodeResponse, getExecutable


# 创建API路由
@app.post("/generate_executable_code", response_model=CodeResponse)
async def generate_code(request: CodeRequest):
    print(request)
    code = getExecutable(request.function_name, request.arguments, request.doc_string)
    print(code)
    return {"code": code}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)