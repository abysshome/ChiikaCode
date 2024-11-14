import csv
import os

import openpyxl
import uvicorn
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document
from langchain.schema.output_parser import StrOutputParser
from langchain.schema.runnable import RunnablePassthrough
from langchain.text_splitter import (CharacterTextSplitter,
                                     RecursiveCharacterTextSplitter)
from langchain_community.chat_models import ChatOllama
from langchain_community.document_loaders import (CSVLoader, PyPDFLoader,
                                                  PythonLoader, TextLoader)
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import Chroma
from openpyxl import load_workbook
from pdfminer.high_level import extract_text
from pydantic import BaseModel


class BaseLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        raise NotImplementedError("This method should be overridden by subclasses")

# 定义各种文件类型的加载器
class PythonLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class PyPDFLoader(BaseLoader):
    def load(self):
        return extract_text(self.file_path)

class TextLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class CSVLoader(BaseLoader):
    def load(self):
        with open(self.file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            return list(reader)

class DocxLoader(BaseLoader):
    def load(self):
        doc = DocxDocument(self.file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

class XlsxLoader(BaseLoader):
    def load(self):
        workbook = load_workbook(filename=self.file_path)
        sheet = workbook.active
        data = []
        for row in sheet.iter_rows(values_only=True):
            data.append(row)
        return data

class CLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class JavaLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

class CPPLoader(BaseLoader):
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

# 数据加载
def load_file(file_path):
    if not os.path.exists(file_path):
        print("文件路径不存在")
        return None

    if file_path.endswith('.py'):
        loader = PythonLoader(file_path)
    elif file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path)
    elif file_path.endswith('.csv'):
        loader = CSVLoader(file_path)
    elif file_path.endswith('.docx'):
        loader = DocxLoader(file_path)
    elif file_path.endswith('.xlsx'):
        loader = XlsxLoader(file_path)
    elif file_path.endswith('.c'):
        loader = CLoader(file_path)
    elif file_path.endswith('.java'):
        loader = JavaLoader(file_path)
    elif file_path.endswith('.cpp'):
        loader = CPPLoader(file_path)
    else:
        print("不支持的文件类型")
        return None
    
    documents = loader.load()
    return documents

# 文档分割
def split_documents(documents):
    if isinstance(documents, str):
        documents = [Document(page_content=documents)]
    elif isinstance(documents, list) and all(isinstance(item, str) for item in documents):
        documents = [Document(page_content=item) for item in documents]
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=65536, chunk_overlap=10)
    documents = text_splitter.split_documents(documents)
    return documents

# 嵌入向量获取
def get_embedding():
    model_name = 'moka-ai/m3e-base'
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': True}
    embedding = HuggingFaceBgeEmbeddings(
        model_name=model_name, 
        model_kwargs=model_kwargs, 
        encode_kwargs=encode_kwargs
    )
    return embedding

# 构建向量数据库
def build_vector_db(documents, embedding):
    persist_directory = 'db'
    db = Chroma.from_documents(
        documents, 
        embedding, 
        persist_directory=persist_directory
    )
    return db

# 处理文件夹
def handle_folder(folder_path):
    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)
        if os.path.isdir(file_path):
            handle_folder(file_path)
        else:
            print(f"处理文件: {file_path}")
            if file_path.endswith(('.py', '.pdf', '.txt', '.csv', '.docx', '.xlsx', '.c', '.java', '.cpp')):
                documents = load_file(file_path)
                if documents:
                    documents = split_documents(documents)
                    embedding = get_embedding()
                    build_vector_db(documents, embedding)
def get_vector_db():
    persist_directory = 'db'
    db = Chroma(persist_directory=persist_directory)
    return db                
def build_retriever(db):
    retriever = db.as_retriever()
    return retriever

def build_rag_chain(retriever):
    template = """
        根据context详细解释有关question的内容，并给出答案。
        Question: {question} Context: {context} Answer: """
    prompt = ChatPromptTemplate.from_template(template)
    llm = ChatOllama(model='llama3.2:latest')
    rag_chain = (
        {"context": retriever, "question": RunnablePassthrough()} |
        prompt |
        llm |
        StrOutputParser()
    )
    return rag_chain
# 主函数
def main(file_path):
    if os.path.isdir(file_path):
        handle_folder(file_path)
    else:
        documents = load_file(file_path)
        if documents:
            documents = split_documents(documents)
            embedding = get_embedding()
            build_vector_db(documents, embedding)
    print("所有文件处理完毕")
    db=get_vector_db()
    retriever = build_retriever(db)
    return build_rag_chain(retriever)

async def chat(rag_chain):
    query = input(">>> Please ask a question: ")
    print(">>> RAG-bot:")
    async for chunk in rag_chain.astream(query): # rag_chain.astream(query)返回一个异步生成器
        print(chunk, end='', flush=True) # chunk表示生成的文本片段，end=''表示不换行，flush=True表示立即输出
    print("\n")

if __name__ == "__main__":

    file_or_folder_path = input("请输入文件或文件夹路径: ")
    rag_chain=main(file_or_folder_path)
